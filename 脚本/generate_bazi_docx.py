#!/usr/bin/env python3
"""将八字简批 Markdown 报告转换为排版精良的 .docx"""

import sys, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_font(run, name='微软雅黑', size=Pt(11), bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_line(doc, text, size=Pt(11), bold=False, color=None, align=None, before=0, after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.space_before = Pt(before)
    p.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.6
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def generate_bazi_report(md_path: str, output_path: str = None):
    md_path = Path(md_path)
    if output_path is None:
        output_path = md_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)

    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')

    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # 解析内容
    report_title = ''
    sections = []  # [{title, body_lines, is_pillar}]
    current_section = None

    for line in lines:
        stripped = line.strip()

        # 文档主标题 (# )
        if stripped.startswith('# ') and not stripped.startswith('## '):
            report_title = stripped[2:].strip()
            continue

        # 二级标题 (## )
        if stripped.startswith('## '):
            if current_section:
                sections.append(current_section)
            current_section = {
                'title': stripped[3:].strip(),
                'body': [],
                'is_code_block': False,
                'code_lines': []
            }
            continue

        # 三级标题 (### ) → 作为子标题
        if stripped.startswith('### '):
            if current_section:
                current_section['body'].append(('h3', stripped[4:].strip()))
            continue

        # 分割线
        if stripped == '---':
            if current_section:
                current_section['body'].append(('hr', ''))
            continue

        # 引用
        if stripped.startswith('> '):
            if current_section:
                current_section['body'].append(('quote', stripped[2:].strip()))
            continue

        # 代码块
        if stripped.startswith('```'):
            if current_section:
                if current_section.get('in_code'):
                    current_section['in_code'] = False
                else:
                    current_section['in_code'] = True
            continue

        if current_section and current_section.get('in_code'):
            current_section['body'].append(('code', stripped))
            continue

        # 加粗文本 **...**
        if stripped.startswith('**') and stripped.endswith('**'):
            if current_section:
                current_section['body'].append(('bold_line', stripped[2:-2].strip()))
            continue

        # 列表项
        if re.match(r'^\d+\.\s', stripped):
            if current_section:
                current_section['body'].append(('list', stripped.strip()))
            continue

        # 空行
        if not stripped:
            if current_section:
                current_section['body'].append(('blank', ''))
            continue

        # 普通正文
        if current_section:
            current_section['body'].append(('text', stripped))

    if current_section:
        sections.append(current_section)

    # === 写入文档 ===

    # 封面标题
    add_line(doc, '八 字 命 理 简 批', size=Pt(18), bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12)

    # 副标题
    if report_title:
        add_line(doc, report_title, size=Pt(11), color=(100, 100, 100),
                 align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=24)

    # 分隔线
    add_line(doc, '─' * 40, size=Pt(8), color=(180, 180, 180),
             align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=12)

    # 逐节输出
    for sec in sections:
        # 节标题
        add_line(doc, sec['title'], size=Pt(13), bold=True,
                 before=16, after=6)

        # 内容
        for item_type, text in sec['body']:
            if item_type == 'blank':
                add_line(doc, '', size=Pt(6), before=0, after=0)

            elif item_type == 'hr':
                add_line(doc, '· · ·', size=Pt(9), color=(160, 160, 160),
                         align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=4)

            elif item_type == 'h3':
                add_line(doc, text, size=Pt(11), bold=True, before=8, after=2)

            elif item_type == 'bold_line':
                add_line(doc, text, size=Pt(11), bold=True, before=2, after=2)

            elif item_type == 'quote':
                add_line(doc, text, size=Pt(10), color=(100, 100, 100), before=2, after=2)

            elif item_type == 'list':
                add_line(doc, text, size=Pt(11), before=0, after=1)

            elif item_type == 'code':
                add_line(doc, text, size=Pt(9), color=(80, 80, 80), before=0, after=0)

            else:
                add_line(doc, text, size=Pt(11), before=0, after=4)

    # 页脚
    add_line(doc, '', size=Pt(6), before=16, after=0)
    add_line(doc, '─' * 40, size=Pt(8), color=(180, 180, 180),
             align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4)
    add_line(doc, '以上分析基于传统命理学，仅供参考。人生选择请结合现实独立判断。', size=Pt(9), color=(140, 140, 140),
             align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2)

    doc.save(str(output_path))
    print(f'[OK] 报告已生成: {output_path}')
    return str(output_path)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python generate_bazi_docx.py <md_path> [output_path]')
        sys.exit(1)
    md_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    generate_bazi_report(md_path, output_path)
